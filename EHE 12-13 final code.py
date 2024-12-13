import os
import time
import pandas as pd
from jobspy import scrape_jobs
from docx import Document
from docx.shared import Inches
from difflib import SequenceMatcher
import matplotlib.pyplot as plt
import logging
import matplotlib
import json
import re
from tensorflow import keras
import tensorflow as tf
from keras.models import Model
from tqdm import tqdm
from transformers import AutoTokenizer, TFAutoModelForTokenClassification, pipeline
from nltk import word_tokenize, pos_tag, ne_chunk
from nltk.tree import Tree
from wordcloud import WordCloud
from textblob import TextBlob


# Set Matplotlib backend to avoid Tkinter issues in headless environments
matplotlib.use("Agg")

# Configure logging with dynamic level control
def configure_logging(level):
    logging.basicConfig(level=level, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s")
    return logging.getLogger("JobScraper")

logger = configure_logging(logging.INFO)

# ----------------------------- User Preferences ----------------------------- #
PREFS_FILE = "user_prefs.json"
OUTPUT_PATH = "output/"

os.makedirs(OUTPUT_PATH, exist_ok=True)

def load_user_preferences():
    if os.path.exists(PREFS_FILE):
        with open(PREFS_FILE, "r") as f:
            return json.load(f)
    return {}

def save_user_preferences(preferences):
    with open(PREFS_FILE, "w") as f:
        json.dump(preferences, f)

def get_unique_file_path(file_path):
    if not os.path.exists(file_path):
        return file_path
    base, extension = os.path.splitext(file_path)
    counter = 1
    while os.path.exists(f"{base}_{counter}{extension}"):
        counter += 1
    return f"{base}_{counter}{extension}"

# ----------------------------- Utility Functions ----------------------------- #
def scrape_zip_recruiter_page(page, model=None):
    if model is None:
        raise ValueError("Model is required for scraping ZipRecruiter.")
    # Proceed with scraping logic

def load_keras_nlp_ner_model():
    """
    Loads a BERT backbone and its preprocessor for use with KerasNLP.
    """
    preprocessor = keras_nlp.models.BertPreprocessor.from_preset("bert_base_uncased")
    model = keras_nlp.models.BertBackbone.from_preset("bert_base_uncased")
    return preprocessor, model

def clean_text(text):
    if not isinstance(text, str):
        return "Not available"
    return ' '.join(text.split()).replace("\n", " ").strip()

def extract_skills_nltk(text):
    try:
        tokens = word_tokenize(text)
        tagged_tokens = pos_tag(tokens)
        named_entities = ne_chunk(tagged_tokens)
        skills = [" ".join(c[0] for c in chunk) for chunk in named_entities if isinstance(chunk, Tree)]
        return skills
    except Exception as e:
        logger.error(f"Error during NLTK skill extraction: {e}")
        return []

def extract_skills_tensorflow(nlp, text):
    """
    Extracts named entities using a TensorFlow-based NER pipeline.
    """
    try:
        if not nlp:
            raise RuntimeError("NER model is not loaded. Ensure TensorFlow and model dependencies are properly configured.")
        ner_results = nlp(text)
        # Extract words labeled as skills (MISC or similar) based on NER model
        skills = [{"entity": entity['entity_group'], "word": entity['word']} for entity in ner_results]
        return skills
    except RuntimeError as re:
        logger.warning(f"Fallback mechanism triggered: {re}")
        return []
    except Exception as e:
        logger.error(f"Unexpected error during TensorFlow skill extraction: {e}")
        return []

def analyze_sentiment(text):
    if not text:
        return "Neutral"
    analysis = TextBlob(text)
    if analysis.sentiment.polarity > 0:
        return "Positive"
    elif analysis.sentiment.polarity < 0:
        return "Negative"
    else:
        return "Neutral"

def generate_histogram_and_skill_list(skills, title, skills_doc):
    if not skills:
        logger.info(f"No skills to process for {title}.")
        return

    skills_df = pd.DataFrame(skills, columns=['Skill'])
    top_n = 20  # Limit to top N skills
    skill_counts = skills_df['Skill'].value_counts().head(top_n)

    plt.figure(figsize=(min(len(skill_counts) * 0.5, 20), 6))
    skill_counts.plot(kind="bar", title=f"Top {top_n} Skills Frequency for {title}")
    plt.xlabel("Skills")
    plt.ylabel("Frequency")
    plt.xticks(rotation=45, ha="right")

    histogram_path = os.path.join(OUTPUT_PATH, f"{title.replace(' ', '_')}_skills_histogram.png")
    plt.tight_layout()
    plt.savefig(histogram_path)
    plt.close()

    skills_doc.add_heading(f"Skills Histogram for {title}", level=2)
    skills_doc.add_picture(histogram_path, width=Inches(6))

    skills_doc.add_heading(f"Skills List for {title}", level=2)
    for skill, count in skill_counts.items():
        skills_doc.add_paragraph(f"{skill}: {count}")
def sanitize_filename(filename):
    """
    Replaces invalid characters in filenames with underscores.
    """
    return re.sub(r'[<>:"/\\|?*]', '_', filename)


def generate_word_cloud(text, title, combined_doc, background_color="white", width=800, height=400):
    if not text:
        logger.info(f"No text available for word cloud generation for {title}.")
        return

    # Sanitize the title to make it a valid filename
    sanitized_title = sanitize_filename(title)
    wordcloud_path = os.path.join(OUTPUT_PATH, f"{sanitized_title}_wordcloud.png")
    os.makedirs(os.path.dirname(wordcloud_path), exist_ok=True)

    # Generate the word cloud
    wordcloud = WordCloud(width=width, height=height, background_color=background_color).generate(text)
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation="bilinear")
    plt.axis("off")
    plt.title(f"Word Cloud for {title}")
    
    # Save the word cloud
    plt.savefig(wordcloud_path)
    plt.close()

    # Add to the document
    combined_doc.add_heading(f"Word Cloud for {title}", level=2)
    combined_doc.add_picture(wordcloud_path, width=Inches(6))

def generate_summary_report(all_skills, combined_doc):
    combined_doc.add_heading("Summary Report", level=1)
    skill_counts = pd.Series(all_skills).value_counts()

    combined_doc.add_paragraph("Top 10 Skills Identified:")
    for skill, count in skill_counts.head(10).items():
        combined_doc.add_paragraph(f"- {skill}: {count}")

def process_job_title(search_term, location, combined_doc, skills_doc, nlp, all_skills, use_keras_nlp=False, preprocessor=None, keras_model=None):
    sites = ["indeed", "linkedin", "zip_recruiter", "glassdoor"]
    all_jobs = pd.DataFrame()

    for site in sites:
        try:
            jobs = scrape_jobs(site_name=site, search_term=search_term, location=location, results_wanted=50)
            if not jobs.empty:
                all_jobs = pd.concat([jobs for jobs in all_jobs if not jobs.empty], ignore_index=True)
        except Exception as e:
            logger.warning(f"Error during scraping {site}: {e}")

    if not all_jobs.empty:
        all_jobs.drop_duplicates(subset=['title', 'company', 'location'], inplace=True)
        combined_doc.add_heading(f"{search_term} Jobs in {location}", level=1)
        combined_doc.add_paragraph(f"Number of Jobs: {len(all_jobs)}")

        for idx, row in all_jobs.iterrows():
            title = clean_text(row.get('title', 'Job title not available'))
            company = clean_text(row.get('company', 'Company not available'))
            job_location = clean_text(row.get('location', 'Location not available'))
            description = clean_text(row.get('description', 'Responsibilities not available'))

            combined_doc.add_heading(f"{idx + 1}. {title}", level=2)
            combined_doc.add_paragraph(f"Company: {company}")
            combined_doc.add_paragraph(f"Location: {job_location}")
            combined_doc.add_paragraph("Description:")
            combined_doc.add_paragraph(description)

            # Choose entity extraction method
            if use_keras_nlp and preprocessor and keras_model:
                entities = extract_skills_keras_nlp(preprocessor, keras_model, description)
            else:
                entities = extract_skills_tensorflow(nlp, description)

            for entity in entities:
                combined_doc.add_paragraph(f"Entity: {entity['token']} ({entity['entity']})")
                all_skills.append(entity['token'])

        generate_histogram_and_skill_list(all_skills, search_term, skills_doc)
    else:
        logger.info(f"No jobs found for {search_term}.")

def extract_skills_keras_nlp(preprocessor, model, text):
    """
    Extracts entities from text using KerasNLP.
    """
    try:
        # Preprocess the text
        inputs = preprocessor(text)
        
        # Get predictions from the model
        outputs = model(inputs)
        logits = outputs["sequence_output"]  # Shape: (batch_size, seq_len, num_labels)
        
        # Map predictions to tokens
        predicted_labels = tf.argmax(logits, axis=-1).numpy().flatten()
        tokens = preprocessor.tokenizer.tokenize(text)
        
        # Match tokens with predicted entities
        entities = []
        for token, label in zip(tokens, predicted_labels):
            if label != 0:  # Assuming 0 corresponds to the 'O' (outside entity) tag
                entities.append({"token": token, "entity": label})
        return entities
    except Exception as e:
        logger.error(f"Error during KerasNLP entity extraction: {e}")
        return []



# ----------------------------- Main Execution ----------------------------- #
if __name__ == "__main__":
    preferences = load_user_preferences()

    logging_level = preferences.get("logging_level", "INFO").upper()
    logger = configure_logging(getattr(logging, logging_level, logging.INFO))

    job_titles = preferences.get("job_titles", [])
    location = preferences.get("location", "")

    # Ask for input if preferences are missing
    if not job_titles or not location:
        job_titles = input("Enter job titles (comma-separated): ").split(",")
        location = input("Enter the location: ")
        preferences["job_titles"] = job_titles
        preferences["location"] = location
        save_user_preferences(preferences)

    # Initialize KerasNLP components
    use_keras_nlp = True
    preprocessor, keras_model, nlp = None, None, None

    if use_keras_nlp:
        try:
            import keras_nlp
            preprocessor, keras_model = load_keras_nlp_ner_model()
            logger.info("KerasNLP model loaded successfully.")
        except Exception as e:
            logger.error(f"Error loading KerasNLP model: {e}")
    else:
        try:
            ner_model_name = "dslim/bert-base-NER"
            ner_tokenizer = AutoTokenizer.from_pretrained(ner_model_name)
            ner_model = TFAutoModelForTokenClassification.from_pretrained(ner_model_name)
            nlp = pipeline("ner", model=ner_model, tokenizer=ner_tokenizer, aggregation_strategy="simple")
            logger.info("Hugging Face NER pipeline loaded successfully.")
        except Exception as e:
            logger.error(f"Error loading Hugging Face pipeline: {e}")

    # Initialize documents
    combined_doc = Document()
    skills_doc = Document()
    all_skills = []

    # Process each job title
    for job_title in job_titles:
        process_job_title(
            job_title.strip(),
            location,
            combined_doc,
            skills_doc,
            nlp,
            all_skills,
            use_keras_nlp=use_keras_nlp,
            preprocessor=preprocessor,
            keras_model=keras_model,
        )

    # Generate summary report
    generate_summary_report(all_skills, combined_doc)

    # Save results
    combined_doc_path = get_unique_file_path(os.path.join(OUTPUT_PATH, "Job_Report.docx"))
    skills_doc_path = get_unique_file_path(os.path.join(OUTPUT_PATH, "Skills_Report.docx"))

    combined_doc.save(combined_doc_path)
    skills_doc.save(skills_doc_path)

    logger.info(f"Reports saved at: {combined_doc_path} and {skills_doc_path}")