import os
import time
import pandas as pd
from jobspy import scrape_jobs
from docx import Document
from docx.shared import Inches
from difflib import SequenceMatcher
import matplotlib.pyplot as plt
import logging
import matplotlib
import json
from tqdm import tqdm

# Set Matplotlib to Agg backend to avoid Tkinter issues
matplotlib.use("Agg")

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s")
logger = logging.getLogger("JobScraper")

# ----------------------------- User Preferences ----------------------------- #
PREFS_FILE = "user_prefs.json"

def load_user_preferences():
    """Load saved user preferences."""
    if os.path.exists(PREFS_FILE):
        with open(PREFS_FILE, "r") as f:
            return json.load(f)
    return {}

def save_user_preferences(preferences):
    """Save user preferences to a JSON file."""
    with open(PREFS_FILE, "w") as f:
        json.dump(preferences, f)

# ----------------------------- TensorFlow and NLTK Setup ----------------------------- #
USE_TENSORFLOW = input("Do you want to enable TensorFlow-based functionalities? (y/n): ").strip().lower() == 'y'
USE_NLTK = input("Do you want to enable NLTK-based functionalities? (y/n): ").strip().lower() == 'y'

if USE_TENSORFLOW:
    try:
        import tensorflow as tf
        from transformers import AutoTokenizer, TFAutoModelForTokenClassification, pipeline

        print("Loading TensorFlow pre-trained NER model...")
        ner_model_name = "dslim/bert-base-NER"
        ner_tokenizer = AutoTokenizer.from_pretrained(ner_model_name)
        ner_model = TFAutoModelForTokenClassification.from_pretrained(ner_model_name)
        nlp = pipeline("ner", model=ner_model, tokenizer=ner_tokenizer, aggregation_strategy="simple")
    except ImportError as e:
        logger.error("TensorFlow or required libraries are not installed. Disabling TensorFlow functionality.")
        USE_TENSORFLOW = False

if USE_NLTK:
    try:
        import nltk
        nltk.download('punkt')
        nltk.download('stopwords')
        nltk.download('averaged_perceptron_tagger')
        nltk.download('wordnet')
        from nltk.corpus import stopwords
        from nltk.tokenize import word_tokenize
        from nltk.stem import WordNetLemmatizer
        from nltk import pos_tag

        lemmatizer = WordNetLemmatizer()
        stop_words = set(stopwords.words('english'))
    except ImportError as e:
        logger.error("NLTK or required libraries are not installed. Disabling NLTK functionality.")
        USE_NLTK = False

# ----------------------------- Utility Functions ----------------------------- #
def clean_text(text):
    """Clean and normalize text."""
    if not isinstance(text, str):  # Handle non-string values
        return "Not available"
    return ' '.join(text.split()).replace("\n", " ").strip()

def calculate_avg_salary(salary):
    """Calculate the average salary from a salary string."""
    if not salary or salary.lower() == "not available":
        return "Not available"
    try:
        salary_range = [float(s.strip('USD').replace(',', '').replace(' ', '').replace('/yearly', '').replace('$', ''))
                        for s in salary.split('-')]
        if len(salary_range) == 2:
            return f"USD{sum(salary_range) / 2:,.2f}/yearly"
        return salary
    except Exception as e:
        logger.warning(f"Error parsing salary: {e}")
        return "Not available"

def generate_histogram_and_skill_list(skills, title, skills_doc):
    """Generates a histogram and skill frequency list dynamically."""
    if not skills:
        logger.info(f"No skills to process for {title}.")
        return

    skills_df = pd.DataFrame(skills, columns=['Skill'])
    skill_counts = skills_df['Skill'].value_counts()

    # Dynamically scale the figure size
    num_skills = len(skill_counts)
    plt.figure(figsize=(max(6, num_skills // 2), 5))
    skill_counts.plot(kind="bar", title=f"Skill Frequency for {title}")
    plt.xlabel("Skills")
    plt.ylabel("Frequency")
    plt.xticks(rotation=45, ha="right")

    # Save and embed the histogram
    histogram_path = f"{title.replace(' ', '_')}_skills_histogram.png"
    plt.tight_layout()
    plt.savefig(histogram_path)
    plt.close()

    # Add histogram to skills document
    skills_doc.add_heading(f"Skills Histogram for {title}", level=2)
    skills_doc.add_picture(histogram_path, width=Inches(6))
    os.remove(histogram_path)

    # Add skill frequency list to skills document
    skills_doc.add_heading(f"Skills List for {title}", level=2)
    for skill, count in skill_counts.items():
        skills_doc.add_paragraph(f"{skill}: {count}")

def get_unique_file_path(file_path):
    """Generate unique file paths to avoid overwriting."""
    if not os.path.exists(file_path):
        return file_path
    base, extension = os.path.splitext(file_path)
    counter = 1
    while os.path.exists(f"{base}_{counter}{extension}"):
        counter += 1
    return f"{base}_{counter}{extension}"

def is_similar(a, b, threshold=0.5):
    """Determine if two strings are similar based on a threshold."""
    return SequenceMatcher(None, a.lower(), b.lower()).ratio() > threshold

# ----------------------------- Job Processing ----------------------------- #
def extract_skills_nltk(text):
    """Extract skills from text using NLTK."""
    if not USE_NLTK:
        return []
    try:
        tokens = word_tokenize(text)
        filtered_tokens = [lemmatizer.lemmatize(w.lower()) for w in tokens if w.isalpha() and w.lower() not in stop_words]
        tagged_tokens = pos_tag(filtered_tokens)
        skills = [word for word, pos in tagged_tokens if pos.startswith('NN')]
        return skills
    except Exception as e:
        logger.error(f"Error during NLTK skill extraction: {e}")
        return []

def extract_skills_tensorflow(text):
    """Extract skills from text using TensorFlow NER."""
    if not USE_TENSORFLOW:
        return []
    try:
        ner_results = nlp(text)
        skills = [entity['word'] for entity in ner_results if entity['entity_group'] == 'MISC']
        return skills
    except Exception as e:
        logger.error(f"Error during TensorFlow skill extraction: {e}")
        return []

def process_job_title(search_term, location, combined_doc, skills_doc, all_skills):
    """Process each job title and extract skills."""
    sites = ["indeed", "linkedin", "zip_recruiter", "glassdoor"]
    all_jobs = pd.DataFrame()

    for site in sites:
        try:
            jobs = scrape_jobs(site_name=site, search_term=search_term, location=location, results_wanted=50)
            if not jobs.empty:
                all_jobs = pd.concat([all_jobs, jobs], ignore_index=True)
        except Exception as e:
            logger.warning(f"Error during scraping {site}: {e}")

    if not all_jobs.empty:
        all_jobs.drop_duplicates(subset=['title', 'company', 'location'], inplace=True)
        combined_doc.add_heading(f"{search_term} Jobs in {location}", level=1)
        combined_doc.add_paragraph(f"Number of Jobs: {len(all_jobs)}\n")

        job_skills = []
        for idx, row in tqdm(all_jobs.iterrows(), total=len(all_jobs), desc=f"Processing {search_term}", unit="job"):
            title = clean_text(row.get('title', 'Job title not available'))
            company = clean_text(row.get('company', 'Company not available'))
            job_location = clean_text(row.get('location', 'Location not available'))
            salary = clean_text(row.get('salary', 'Not available'))
            avg_salary = calculate_avg_salary(salary)
            responsibilities = clean_text(row.get('description', row.get('job_function', row.get('company_description', 'Responsibilities not available.'))))
            responsibilities_list = responsibilities.split('. ') if responsibilities else ['Responsibilities not available.']

            # Add job details to the main document
            combined_doc.add_heading(f'{idx + 1}. {title}', level=2)
            combined_doc.add_paragraph(f'Company: {company}')
            combined_doc.add_paragraph(f'Location: {job_location}')
            combined_doc.add_paragraph(f'Salary: {salary}')
            combined_doc.add_paragraph(f'Average Salary: {avg_salary}')
            combined_doc.add_paragraph('Key Responsibilities:')
            for responsibility in responsibilities_list:
                if responsibility.strip():
                    combined_doc.add_paragraph(responsibility.strip())

            # Extract skills using NLTK and TensorFlow
            skills_nltk = extract_skills_nltk(responsibilities)
            skills_tf = extract_skills_tensorflow(responsibilities)
            skills = list(set(skills_nltk + skills_tf))
            if not skills:  # If no skills extracted, use predefined skills
                skills = [
                    'Strong knowledge of OSHA, state, and federal regulations',
                    'Experience with LOTO, confined spaces, and PSM',
                    'Knowledge of wastewater and stormwater regulations',
                    'Strong communication, leadership, and problem-solving skills',
                    'Proficiency with Microsoft Office Suite'
                ]
            job_skills.extend(skills)
            all_skills.extend(skills)

        # Generate and add a skills histogram and list format to both documents
        if job_skills:
            skills_df = pd.DataFrame(job_skills, columns=['Skill'])
            skill_counts = skills_df['Skill'].value_counts()

            # Add skills list in the main document
            combined_doc.add_paragraph("Skills List:")
            for skill, count in skill_counts.items():
                combined_doc.add_paragraph(f"{skill} ({count})")

            # Generate histogram and skill list for this title
            generate_histogram_and_skill_list(job_skills, search_term, skills_doc)
    else:
        print(f"No similar jobs found for '{search_term}'.")

    return all_jobs, all_skills

# ----------------------------- Main Execution ----------------------------- #
if __name__ == "__main__":
    preferences = load_user_preferences()

    reuse_titles = preferences.get("job_titles", [])
    reuse_location = preferences.get("location", "")

    if reuse_titles and reuse_location:
        reuse = input("Do you want to reuse the saved job titles and location? (y/n): ").strip().lower() == 'y'
        if reuse:
            job_titles = reuse_titles
            location = reuse_location
        else:
            job_titles_input = input("Enter job titles (comma-separated): ")
            location = input("Enter the location: ")
            job_titles = [title.strip() for title in job_titles_input.split(",") if title.strip()]
            preferences["job_titles"] = job_titles
            preferences["location"] = location
    else:
        job_titles_input = input("Enter job titles (comma-separated): ")
        location = input("Enter the location: ")
        job_titles = [title.strip() for title in job_titles_input.split(",") if title.strip()]
        preferences["job_titles"] = job_titles
        preferences["location"] = location

    save_user_preferences(preferences)

    if not job_titles:
        print("No job titles provided.")
        exit()

    combined_doc = Document()
    skills_doc = Document()
    all_skills = []

    start_time = time.time()
    for idx, job in enumerate(job_titles):
        elapsed_time = time.time() - start_time
        avg_time = elapsed_time / (idx + 1)
        remaining_time = avg_time * (len(job_titles) - idx - 1)
        print(f"Processing '{job}'... Estimated time left: {int(remaining_time // 60)} minutes and {int(remaining_time % 60)} seconds")

        _, job_skills = process_job_title(job, location, combined_doc, skills_doc, all_skills)

    # Generate overall skills histogram and list
    if all_skills:
        generate_histogram_and_skill_list(all_skills, "All Jobs", skills_doc)

    downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
    combined_doc_path = get_unique_file_path(os.path.join(downloads_folder, "Job_Report.docx"))
    skills_doc_path = get_unique_file_path(os.path.join(downloads_folder, "Skills_Report.docx"))
    combined_doc.save(combined_doc_path)
    skills_doc.save(skills_doc_path)

    print(f"Reports saved: {combined_doc_path} and {skills_doc_path}")
